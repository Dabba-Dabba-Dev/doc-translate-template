import os
import json
from PIL import Image, ImageOps
from layout_combined_enhanced import EnhancedDocumentReconstructor
from file_utils import pdf_to_images
from ocr_engine import extract_blocks_with_boxes
import requests
from flask import send_file
from flask import Flask, request, jsonify
import shutil
import re

def looks_non_translatable(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True  # skip empty

    # Emails, URLs
    if re.search(r'\S+@\S+|\bhttps?://\S+', stripped):
        return True

    # Phone numbers
    if re.match(r'^\+?\d[\d\s().-]{5,}$', stripped):
        return True

    # Dates (many formats)
    if re.match(r'^(\d{1,2}(st|nd|rd|th)?\s+[A-Za-z]+,?\s*\d{0,4}|[A-Za-z]+\s+\d{1,2},?\s*\d{0,4}|\d{4}-\d{2}-\d{2})$', stripped):
        return True

    # Pure numbers / decimals / comma numbers
    if re.match(r'^[\d,.]+$', stripped):
        return True

    words = stripped.split()

    # Acronyms / org names (all caps or mostly caps)
    if all(c.isupper() or not c.isalpha() for c in stripped) and len(stripped) > 1:
        return True

    # Short proper nouns / titles
    if len(words) <= 3 and all(w[0].isupper() for w in words if w):
        return True

    # Single punctuation / symbols only
    if re.match(r'^[^\w\s]+$', stripped):
        return True

    return False

def clean_directory(path):
    """Delete all contents of a directory but keep the directory itself."""
    if os.path.exists(path):
        for filename in os.listdir(path):
            file_path = os.path.join(path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")
    else:
        os.makedirs(path, exist_ok=True)

app = Flask(__name__)
def translate_line_remote(text, src_lang, tgt_lang):
    try:
        resp = requests.post(
            "http://translator:5001/translate",
            json={"text": text, "src_lang": src_lang, "tgt_lang": tgt_lang}
        )
        resp.raise_for_status()
        return resp.json().get("translated", "[Translation Failed]")
    except Exception as e:
        print("Translation API call failed:", e)
        return "[Translation Failed]"
@app.route("/process", methods=["POST"])
def process_document():
    global INPUT_PDF, TEMP_IMAGE_DIR, OCR_OUTPUT_DIR, TRANSLATION_OUTPUT_DIR, LAYOUT_OUTPUT_DIR, FINAL_PDF_OUTPUT, FINAL_DOCX_OUTPUT

    # Define paths first
    INPUT_PDF = "/tmp/input.pdf"
    TEMP_IMAGE_DIR = "debug_results/temp_images"
    OCR_OUTPUT_DIR = "debug_results/ocr_results"
    TRANSLATION_OUTPUT_DIR = "debug_results/translated_results"
    LAYOUT_OUTPUT_DIR = "debug_results/layout_output"
    FINAL_PDF_OUTPUT = "data/output_docs/final_output.pdf"
    FINAL_DOCX_OUTPUT = "data/output_docs/final_output.docx"
    # Make sure directories exist and are clean
    clean_directory(TEMP_IMAGE_DIR)
    clean_directory(OCR_OUTPUT_DIR)
    clean_directory(TRANSLATION_OUTPUT_DIR)
    clean_directory(LAYOUT_OUTPUT_DIR)
    clean_directory(os.path.dirname(FINAL_PDF_OUTPUT))
    clean_directory(os.path.dirname(FINAL_DOCX_OUTPUT))
    os.makedirs(TEMP_IMAGE_DIR, exist_ok=True)
    os.makedirs(OCR_OUTPUT_DIR, exist_ok=True)
    os.makedirs(TRANSLATION_OUTPUT_DIR, exist_ok=True)
    os.makedirs(LAYOUT_OUTPUT_DIR, exist_ok=True)
    os.makedirs(os.path.dirname(FINAL_PDF_OUTPUT), exist_ok=True)
    os.makedirs(os.path.dirname(FINAL_DOCX_OUTPUT), exist_ok=True)

    file = request.files.get("file")
    src_lang = request.form.get("src_lang", "pl_PL")
    tgt_lang = request.form.get("tgt_lang", "en_XX")
    
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    file.save(INPUT_PDF)
    ext = os.path.splitext(file.filename)[1].lower()

    if ext == ".pdf":
        image_paths = step1_pdf_to_images()
    else:
        from PIL import Image, ImageOps
        img = Image.open(INPUT_PDF)
        img = ImageOps.exif_transpose(img)
        image_path = os.path.join(TEMP_IMAGE_DIR, "input.jpg")
        img.convert("RGB").save(image_path, "JPEG")
        image_paths = [image_path]
    if not image_paths:
        return jsonify({
            "message": "No pages to process",
            "download_url": None
        }), 400

    # Run pipeline
    step2_apply_ocr(image_paths, src_lang)
    step3_translate_with_lang(src_lang, tgt_lang)
    step4_reconstruct_layout()
    step5_create_docx()

    return jsonify({"message": "Processing complete", "download_url": "/download-final-docx"})






def step1_pdf_to_images():
    images = pdf_to_images(INPUT_PDF)
    image_paths = []
    for idx, image in enumerate(images):
        image_filename = f"page_{idx + 1}.jpg"  
        image_path = os.path.join(TEMP_IMAGE_DIR, image_filename)
        image.save(image_path)
        image_paths.append(image_path)
        print(f"[Step 1] Saved page {idx + 1} to {image_path}")
    return image_paths

def step2_apply_ocr(image_paths, src_lang):
    os.makedirs(OCR_OUTPUT_DIR, exist_ok=True)
    
    for image_path in image_paths:
        try:
            image = Image.open(image_path)
            image = ImageOps.exif_transpose(image)

            base_name = os.path.splitext(os.path.basename(image_path))[0]
            overlay_path = os.path.join(OCR_OUTPUT_DIR, f"{base_name}_overlay.png")
            json_path = os.path.join(OCR_OUTPUT_DIR, f"{base_name}.json")

            ocr_results = extract_blocks_with_boxes(image, image_path=overlay_path)

            with open(json_path, "w", encoding='utf-8') as f:
                json.dump(ocr_results, f, ensure_ascii=False, indent=2)

            print(f"[Step 2] OCR complete: {json_path}")

        except Exception as e:
            print(f"[ERROR] OCR failed for {image_path}: {e}")

def step3_translate_with_lang(src_lang, tgt_lang):
    print(f"[Step 3] Starting translation: src={src_lang}, tgt={tgt_lang}")

    os.makedirs(TRANSLATION_OUTPUT_DIR, exist_ok=True)
    json_files = [f for f in os.listdir(OCR_OUTPUT_DIR) if f.endswith(".json")]

    for file_name in json_files:
        input_path = os.path.join(OCR_OUTPUT_DIR, file_name)
        output_path = os.path.join(TRANSLATION_OUTPUT_DIR, file_name)

        with open(input_path, "r", encoding="utf-8") as f:
            ocr_data = json.load(f)

        texts = [item["block_text"] for item in ocr_data if item.get("block_text")]
        translated_texts = []

        for idx, text in enumerate(texts):
            print(f"[Step 3] Translating block {idx + 1}/{len(texts)}...")
            lines = text.split("\n")
            translated_lines = []
            for line in lines:
                if looks_non_translatable(line):
                    translated_lines.append(line)  # keep original
                else:
                    translated_lines.append(translate_line_remote(line, src_lang, tgt_lang))
            translated_block = "\n".join(translated_lines)
            translated_texts.append({"original": text, "translated": translated_block})

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(translated_texts, f, ensure_ascii=False, indent=2)

        print(f"[Step 3] Translation saved: {output_path}")

def step4_reconstruct_layout():
    reconstructor = EnhancedDocumentReconstructor()
    os.makedirs(LAYOUT_OUTPUT_DIR, exist_ok=True)
    
    ocr_files = [f for f in os.listdir(OCR_OUTPUT_DIR) if f.endswith(".json")]

    for ocr_file in ocr_files:
        ocr_path = os.path.join(OCR_OUTPUT_DIR, ocr_file)
        translation_path = os.path.join(TRANSLATION_OUTPUT_DIR, ocr_file)
        base_name = os.path.splitext(ocr_file)[0]

        if not os.path.exists(translation_path):
            print(f"[Step 4] Missing translation file for {ocr_file}")
            continue

        try:
            ocr_data, translation_data = reconstructor.load_data(ocr_path, translation_path)
            
            # Get structured content instead of generating image
            structured_content = reconstructor.get_structured_content(ocr_data, translation_data)
            
            # Save structured content for DOCX generation
            content_path = os.path.join(LAYOUT_OUTPUT_DIR, f"{base_name}_content.json")
            with open(content_path, 'w', encoding='utf-8') as f:
                json.dump(structured_content, f, ensure_ascii=False, indent=2)

            print(f"[Step 4] Structured content saved: {content_path}")

        except Exception as e:
            print(f"[ERROR] Content structuring failed for {ocr_file}: {e}")
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document

def step5_create_docx():
    global FINAL_DOCX_OUTPUT
    
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    if not os.path.exists(LAYOUT_OUTPUT_DIR):
        print(f"[Step 5] No layout directory found at {LAYOUT_OUTPUT_DIR}")
        return

    content_files = sorted([
        os.path.join(LAYOUT_OUTPUT_DIR, f)
        for f in os.listdir(LAYOUT_OUTPUT_DIR)
        if f.endswith("_content.json") and os.path.isfile(os.path.join(LAYOUT_OUTPUT_DIR, f))
    ])

    if not content_files:
        print(f"[Step 5] No content files found in {LAYOUT_OUTPUT_DIR}")
        return

    for page_idx, content_file in enumerate(content_files):
        if page_idx > 0:
            doc.add_page_break()
            
        with open(content_file, 'r', encoding='utf-8') as f:
            structured_content = json.load(f)
        
        last_y = 0
        for block in structured_content:
            text = block['text']
            is_bold = block['is_bold']
            alignment = block['alignment']
            current_y = block['original_y']
            
            if current_y - last_y > 0.05: 
                doc.add_paragraph()  
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            
            run.font.size = Pt(12)  
            run.font.bold = is_bold
            run.font.name = 'Times New Roman'
            
            if alignment == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            last_y = current_y

    os.makedirs(os.path.dirname(FINAL_DOCX_OUTPUT), exist_ok=True)
    
    try:
        doc.save(FINAL_DOCX_OUTPUT)
        print(f"[Step 5] Final DOCX saved to {FINAL_DOCX_OUTPUT}")
    except PermissionError:
        FINAL_DOCX_OUTPUT = "/tmp/final_output.docx"
        doc.save(FINAL_DOCX_OUTPUT)
        print(f"[Step 5] Final DOCX saved to {FINAL_DOCX_OUTPUT} (fallback location)")
@app.route("/download-final-docx", methods=["GET"])
def download_final_docx():
    try:
        return send_file(
            FINAL_DOCX_OUTPUT,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='translated_document.docx'
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 404

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
